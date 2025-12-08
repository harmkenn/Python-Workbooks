import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime, timedelta
import re
import urllib.parse

st.title("Blogger → DOCX Exporter (Web Crawl, All Posts)")

st.markdown("""
- Paste the Blogger blog URL (e.g., micheleharmon13.blogspot.com).  
- Choose start and end dates.  
- The app will crawl pages by date and save posts as separate DOCX files.
""")

blog_url = st.text_input("Blogger URL (without https://)")
start_date = st.date_input("Start Date")
end_date = st.date_input("End Date")

if st.button("Export Posts to DOCX"):
    if not blog_url:
        st.warning("Please enter a valid Blogger URL.")
    elif start_date > end_date:
        st.warning("Start date must be before end date.")
    else:
        st.info("Starting crawl...")
        current_date = end_date + timedelta(days=1)  # start just after end date
        saved_files = []
        progress_text = st.empty()
        progress_bar = st.progress(0)
        total_posts = 0

        while current_date > start_date:
            # Construct the search URL
            updated_max = current_date.strftime("%Y-%m-%dT23:59:59-06:00")
            search_url = f"https://{blog_url}/search?updated-max={updated_max}&max-results=7"
            response = requests.get(search_url)
            if response.status_code != 200:
                st.warning(f"Failed to fetch: {search_url}")
                break

            soup = BeautifulSoup(response.text, "html.parser")
            posts = soup.find_all("h3", class_="post-title")  # adjust selector if needed
            if not posts:
                st.info("No more posts found.")
                break

            for post in posts:
                link_tag = post.find("a")
                if not link_tag:
                    continue
                post_url = link_tag['href']
                post_title = link_tag.get_text(strip=True)

                # Fetch post page
                post_resp = requests.get(post_url)
                post_soup = BeautifulSoup(post_resp.text, "html.parser")

                # Extract content (adjust the div class to your blog)
                content_div = post_soup.find("div", class_="post-body")
                if not content_div:
                    continue

                doc = Document()
                doc.add_heading(post_title, level=1)

                # Loop through content
                for elem in content_div.descendants:
                    if elem.name == "img":
                        img_url = elem.get("src")
                        if img_url:
                            try:
                                img_data = requests.get(img_url).content
                                doc.add_picture(BytesIO(img_data), width=Inches(5))
                            except:
                                pass
                    elif isinstance(elem, str):
                        text = elem.strip()
                        if text:
                            doc.add_paragraph(text)
                    elif elem.name == "br":
                        doc.add_paragraph("")

                # Use date from post if available
                date_tag = post_soup.find("h2", class_="date-header")
                if date_tag:
                    try:
                        pub_date = datetime.strptime(date_tag.get_text(strip=True), "%A, %B %d, %Y").date()
                        date_str = pub_date.strftime("%Y-%m-%d")
                    except:
                        date_str = "unknown_date"
                else:
                    date_str = "unknown_date"

                safe_title = re.sub(r"[^a-zA-Z0-9_\-]", "_", post_title)
                filename = f"{date_str}_{safe_title}.docx"
                doc.save(filename)
                saved_files.append(filename)
                total_posts += 1

                progress_text.text(f"Saved {total_posts}: {filename}")
                progress_bar.progress(min(total_posts / 100, 1.0))  # rough progress

            # Move current_date back by 7 days (or adjust depending on max-results)
            current_date -= timedelta(days=7)

        if saved_files:
            st.success(f"Saved {len(saved_files)} posts as DOCX:")
            for f in saved_files:
                st.write(f"- {f}")
        else:
            st.warning("No posts were saved.")
