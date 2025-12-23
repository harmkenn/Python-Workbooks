import streamlit as st
import requests
import feedparser
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO
import re

st.title("Blogspot to Word Exporter (With Images)")

blog_url = st.text_input("Enter your Blogspot URL (e.g., micheleharmon13.blogspot.com):")

def clean_url(url):
    url = url.strip()
    url = url.replace("https://", "").replace("http://", "").rstrip("/")
    return url

def fetch_feed(url):
    feed_url = f"https://{url}/feeds/posts/default?alt=json&max-results=999"
    return feed_url

def extract_images_and_text(html_content):
    soup = BeautifulSoup(html_content, "html.parser")

    # Convert <br> and <p> to newlines
    for br in soup.find_all("br"):
        br.replace_with("\n")
    for p in soup.find_all("p"):
        p.insert_after("\n")

    text = soup.get_text(separator="\n")

    # Extract image URLs
    images = []
    for img in soup.find_all("img"):
        if "src" in img.attrs:
            images.append(img["src"])

    return text.strip(), images

def download_image(url):
    try:
        return requests.get(url, timeout=10).content
    except:
        return None

if st.button("Download All Posts"):
    if not blog_url:
        st.error("Please enter a Blogspot URL.")
        st.stop()

    cleaned = clean_url(blog_url)
    feed_api_url = fetch_feed(cleaned)

    st.write(f"Fetching posts from: {feed_api_url}")

    response = requests.get(feed_api_url)

    if response.status_code != 200:
        st.error("Couldn't access feed. Blog might block indexing.")
        st.stop()

    data = response.json()

    if "feed" not in data or "entry" not in data["feed"]:
        st.error("No posts found. Blog might be private.")
        st.stop()

    entries = data["feed"]["entry"]
    st.success(f"Found {len(entries)} posts!")

    for idx, entry in enumerate(entries):
        title = entry["title"]["$t"].strip() or f"Post_{idx+1}"

        # Extract HTML from content
        html = entry["content"]["$t"]
        text, images = extract_images_and_text(html)

        # Create Word document
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)

        # Insert images
        for img_url in images:
            img_data = download_image(img_url)
            if img_data:
                doc.add_picture(BytesIO(img_data), width=Inches(5))

        # Safe filename
        safe_title = re.sub(r"[^a-zA-Z0-9_\- ]", "", title)[:50]
        filename = f"{safe_title}.docx"

        doc.save(filename)

    st.success("All posts downloaded as Word documents!")
    st.info("Check the folder where you ran Streamlit — all .docx files are there.")
